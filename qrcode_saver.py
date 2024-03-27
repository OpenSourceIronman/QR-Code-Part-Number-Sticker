import os
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT, WD_PAPER_SIZE

document = Document()
section = document.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_height = WD_PAPER_SIZE.A4.height
section.page_width = WD_PAPER_SIZE.A4.width

# Open the folder containing your imagess
path_to_images = 'images/'

# Create a table with 4 columns
table = document.add_table(rows=1, cols=4)

# Loop through all the image files in the folder for sticker template at avery.com/templates searching for 22816
for i, filename in enumerate(os.listdir(path_to_images)):
    if filename.endswith('.jpg') or filename.endswith('.png'):
        # Add a new row to the table every 4 images
        if i % 4 == 0:
            row_cells = table.add_row().cells

        # Add the image to the current row
        cell = row_cells[i % 4]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(f'{path_to_images}/{filename}', width=Inches(2))

# Save the document
document.save('Avery_QR_Document.docx')



# TODO DO AGAIN FOR 2ND STICKER TEMPLARE

# Create a table with 4 columns
table = document.add_table(rows=1, cols=7)

# Loop through all the image files in the folder for sticker template at chromalabel.com/templates searching for CAL04460
for i, filename in enumerate(os.listdir(path_to_images)):
    if filename.endswith('.jpg') or filename.endswith('.png'):
        # Add a new row to the table every 7 images
        if i % 7 == 0:
            row_cells = table.add_row().cells

        # Add the image to the current row
        cell = row_cells[i % 7]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(f'{path_to_images}/{filename}', width=Inches(2))

# Save the document
document.save('ChromaLabel_QR_Document.docx')
